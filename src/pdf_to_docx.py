#!/usr/bin/env python3
"""
Convert a PDF to an editable DOCX while preserving formatting.

Uses pdf2docx for layout/table extraction, then post-processes the result:
  * Replaces the legacy "Symbol" font (which pdf2docx assigns to bullet glyphs)
    with the actual font used in the PDF (e.g. "SymbolMT") so that bullets
    render correctly in Word instead of as missing-glyph squares.
  * Re-injects the header row of the first table on page 1. pdf2docx
    sometimes drops a table's header row when its column structure differs
    from the data rows below it (e.g. an empty leftmost cell + merged cells).
    The header is re-extracted directly from the PDF via PyMuPDF and
    reinserted with matching shading + bold text.

Usage:
    python pdf_to_docx.py input.pdf [output.docx]

If output is omitted, the .docx file is written next to the PDF with the same
basename. The script is safe to re-run; it overwrites the output.

Tip: pdf2docx's "stream table" heuristic sometimes invents tables out of
aligned text (e.g. label/value blocks). Disable it via --no-stream-tables when
your PDF has key-value lists that should stay as paragraphs.
"""

from __future__ import annotations

import argparse
import shutil
import subprocess
import sys
import tempfile
from collections import Counter
from pathlib import Path
from typing import Callable, Optional

import fitz  # PyMuPDF, already a dependency of pdf2docx
from pdf2docx import Converter
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL

from _version import __author__, __codename__, __release_date__, __version__
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

# Fonts whose names cause Word to render U+2022 (bullet) as a missing-glyph
# square. The legacy "Symbol" PostScript font remaps Unicode code points, so
# the actual bullet character ends up unmapped. Replace with the OpenType
# variant that keeps Unicode mappings intact.
BAD_BULLET_FONTS = {"Symbol"}
REPLACEMENT_BULLET_FONT = "SymbolMT"


# ----- OCR support ---------------------------------------------------------
#
# pdf2docx can't process scanned PDFs at all — its layout analyzer relies on
# real font glyphs with measurable widths, and skips image-only pages with
# the warning "Words count: 0. It might be a scanned pdf, which is not
# supported yet". So for scanned PDFs we bypass pdf2docx entirely: render
# each page, run tesseract directly to get word-level text + positions
# (TSV output), and build a DOCX from that text via python-docx.
#
# Tesseract is NOT bundled — the user must install it themselves (e.g.
# `brew install tesseract tesseract-lang` on macOS, the UB Mannheim
# installer on Windows).


def list_tesseract_languages() -> Optional[list[str]]:
    """Return a sorted list of language codes installed in the local
    tesseract, or ``None`` if tesseract is not on PATH (or otherwise fails
    to respond).

    The output of ``tesseract --list-langs`` looks like::

        List of available languages (3):
        eng
        osd
        slk

    ``osd`` is the orientation/script detection model, not a real language,
    so it is filtered out.
    """
    if shutil.which("tesseract") is None:
        return None
    try:
        result = subprocess.run(
            ["tesseract", "--list-langs"],
            capture_output=True,
            text=True,
            timeout=5,
        )
    except (subprocess.TimeoutExpired, OSError):
        return None

    # Different tesseract versions print to stdout vs stderr; combine both.
    text = (result.stdout or "") + "\n" + (result.stderr or "")
    langs: set[str] = set()
    for line in text.splitlines():
        line = line.strip()
        if not line or ":" in line or " " in line:
            continue
        if all(c.isalnum() or c == "_" for c in line):
            langs.add(line)
    langs.discard("osd")
    return sorted(langs) if langs else None


ProgressCallback = Callable[[int, int], None]


def _merge_lines_to_paragraphs(
    line_records: list[dict],
    dpi: int,
    page_width_px: int,
) -> list[dict]:
    """Greedy-merge sorted line records into paragraph dicts.

    Same merge rules as the original linear parser, factored out so it
    can be applied per-block in the layout-aware path. Lines belong to
    the same paragraph when they share the same tesseract block, have
    similar height (within ±25 %), and sit within ~0.7 line-heights of
    each other vertically. Returns paragraph dicts with ``text``,
    ``font_size_pt``, ``alignment`` and ``space_before_pt``.
    """
    paragraphs: list[dict] = []
    current: list[dict] | None = None
    prev_para_bottom: int | None = None

    def _finalize(group: list[dict], gap_to_prev: int | None) -> dict:
        text = " ".join(ln["text"] for ln in group)
        heights = sorted(ln["height"] for ln in group)
        median_h = heights[len(heights) // 2]
        font_size_pt = max(6.0, min(round(median_h * 72.0 / dpi, 1), 72.0))

        alignment = "left"
        if page_width_px > 0:
            min_left = min(ln["left"] for ln in group)
            max_right = max(ln["right"] for ln in group)
            left_margin = min_left
            right_margin = page_width_px - max_right
            if left_margin > page_width_px * 0.10:
                imbalance = abs(left_margin - right_margin) / page_width_px
                if imbalance < 0.05:
                    alignment = "center"

        space_before_pt = 0.0
        if gap_to_prev is not None and gap_to_prev > median_h * 1.0:
            extra_px = min(gap_to_prev - int(median_h), int(median_h * 4))
            space_before_pt = round(extra_px * 72.0 / dpi, 1)

        return {
            "text": text,
            "font_size_pt": font_size_pt,
            "alignment": alignment,
            "space_before_pt": space_before_pt,
        }

    for line in line_records:
        if current is None:
            current = [line]
            continue

        prev = current[-1]
        same_block = line["block"] == prev["block"]
        ratio = line["height"] / max(prev["height"], 1)
        similar_height = 0.75 <= ratio <= 1.33
        gap = line["top"] - prev["bottom"]
        small_gap = gap < line["height"] * 0.7

        if same_block and similar_height and small_gap:
            current.append(line)
        else:
            paragraphs.append(_finalize(
                current,
                gap_to_prev=(current[0]["top"] - prev_para_bottom)
                            if prev_para_bottom is not None else None,
            ))
            prev_para_bottom = current[-1]["bottom"]
            current = [line]

    if current is not None:
        paragraphs.append(_finalize(
            current,
            gap_to_prev=(current[0]["top"] - prev_para_bottom)
                        if prev_para_bottom is not None else None,
        ))

    return _merge_code_paragraphs(paragraphs)


def _parse_tesseract_tsv_layout(tsv_path: Path, dpi: int = 300) -> dict:
    """Layout-aware parse of tesseract TSV output.

    Returns a structured page layout::

        {
          "page_width_px":  int,           # rendered page bitmap width
          "page_height_px": int,           # rendered page bitmap height
          "blocks": [
              {
                "block_num": int,
                "bbox": (left, top, right, bottom),  # tesseract pixel coords
                "lines": [<line records>],           # for table detection
                "paragraphs": [<paragraph dicts>],   # post line→para merge
              },
              ...
          ],
          "word_boxes": [(left, top, right, bottom), ...],  # for image detection
        }

    The blocks correspond directly to tesseract's level-2 layout units —
    these are the natural starting point for column detection because
    tesseract already groups physically-adjacent runs of text into the
    same block.

    Lines from each block are merged into paragraphs by
    ``_merge_lines_to_paragraphs``; the heading classifier and code
    merger then operate on the paragraph dicts unchanged.
    """
    import csv

    page_width = 0
    page_height = 0
    block_bboxes: dict[int, tuple[int, int, int, int]] = {}
    # (block, par, line) → list of word dicts. Tesseract emits one TSV
    # per page so page_num is always 1 here; we drop it from the key.
    words_by_line: dict[tuple[int, int, int], list[dict]] = {}
    word_boxes: list[tuple[int, int, int, int]] = []

    with open(tsv_path, encoding="utf-8") as f:
        reader = csv.reader(f, delimiter="\t", quoting=csv.QUOTE_NONE)
        header = next(reader, None)
        if header is None:
            return {"page_width_px": 0, "page_height_px": 0,
                    "blocks": [], "word_boxes": []}
        col = {name: i for i, name in enumerate(header)}
        required = {
            "level", "page_num", "block_num", "par_num", "line_num",
            "word_num", "left", "top", "width", "height", "text",
        }
        if not required.issubset(col):
            return {"page_width_px": 0, "page_height_px": 0,
                    "blocks": [], "word_boxes": []}

        for row in reader:
            if len(row) < len(header):
                continue
            try:
                level = int(row[col["level"]])
                left = int(row[col["left"]])
                top = int(row[col["top"]])
                width = int(row[col["width"]])
                height = int(row[col["height"]])
            except ValueError:
                continue

            if level == 1:  # page
                page_width = width
                page_height = height
                continue

            if level == 2:  # block
                try:
                    block_num = int(row[col["block_num"]])
                except ValueError:
                    continue
                block_bboxes[block_num] = (left, top, left + width, top + height)
                continue

            if level != 5:  # words only
                continue

            try:
                block = int(row[col["block_num"]])
                par = int(row[col["par_num"]])
                line = int(row[col["line_num"]])
                word_num = int(row[col["word_num"]])
            except ValueError:
                continue

            text = row[col["text"]]
            if not text or not text.strip():
                continue

            word_boxes.append((left, top, left + width, top + height))
            words_by_line.setdefault((block, par, line), []).append({
                "word_num": word_num,
                "left": left,
                "top": top,
                "width": width,
                "height": height,
                "text": text,
            })

    # Build line records grouped by block, in tesseract's natural order.
    lines_by_block: dict[int, list[dict]] = {}
    for key in sorted(words_by_line.keys()):
        block, par, line_num = key
        words = sorted(words_by_line[key], key=lambda w: w["word_num"])
        if not words:
            continue
        text = " ".join(w["text"] for w in words)
        heights = sorted(w["height"] for w in words)
        median_h = heights[len(heights) // 2]
        lines_by_block.setdefault(block, []).append({
            "block": block,
            "par": par,
            "line_num": line_num,
            "text": text,
            "top": min(w["top"] for w in words),
            "bottom": max(w["top"] + w["height"] for w in words),
            "left": min(w["left"] for w in words),
            "right": max(w["left"] + w["width"] for w in words),
            "height": median_h,
            "words": words,  # kept for table-cell detection
        })

    blocks_out: list[dict] = []
    for block_num in sorted(lines_by_block.keys()):
        block_lines = lines_by_block[block_num]
        if not block_lines:
            continue
        block_paragraphs = _merge_lines_to_paragraphs(
            block_lines, dpi, page_width,
        )
        bbox = block_bboxes.get(block_num)
        if bbox is None:
            bbox = (
                min(ln["left"] for ln in block_lines),
                min(ln["top"] for ln in block_lines),
                max(ln["right"] for ln in block_lines),
                max(ln["bottom"] for ln in block_lines),
            )
        blocks_out.append({
            "block_num": block_num,
            "bbox": bbox,
            "lines": block_lines,
            "paragraphs": block_paragraphs,
        })

    return {
        "page_width_px": page_width,
        "page_height_px": page_height,
        "blocks": blocks_out,
        "word_boxes": word_boxes,
    }


def _parse_tesseract_tsv(tsv_path: Path, dpi: int = 300) -> list[dict]:
    """Legacy linear parse — flat paragraph list across all blocks.

    Kept as a thin wrapper around the layout-aware parser for any
    callers that don't need the per-block geometry. Internally just
    flattens the layout's blocks back into a single paragraph stream.
    """
    layout = _parse_tesseract_tsv_layout(tsv_path, dpi=dpi)
    out: list[dict] = []
    for block in layout["blocks"]:
        out.extend(block["paragraphs"])
    return out


# ---- Layout region detection (multi-column / image / table) ---------------
#
# After parsing tesseract's TSV into a list of blocks (each with a bbox and
# its own paragraphs), we group them into a higher-level "region" stream
# that the DOCX renderer can walk in reading order. Three region types:
#
#   {"type": "text",    "bbox", "paragraphs"}        — single block of prose
#   {"type": "columns", "bbox", "columns": [...]}    — N side-by-side text
#                                                     blocks; rendered as a
#                                                     borderless N-column
#                                                     table in DOCX
#   {"type": "image",   "bbox", "png_bytes"}         — non-text region cropped
#                                                     out of the page bitmap
#
# Why a borderless table for columns? Word doesn't expose mid-page section
# column changes through python-docx in any clean way; the standard
# workaround is a 1-row N-column borderless table. The visual result is
# identical and it survives round-trips through Word/LibreOffice.


def _detect_table_in_block(block: dict, dpi: int) -> dict | None:
    """Detect a tabular structure inside a single tesseract block.

    Algorithm — *vertical white-corridor detection*:

    1. Strip obvious OCR border noise (single non-alphanumeric chars
       like ``—``, ``|``, ``_`` and ultra-narrow glyphs ≤ 0.1 inch).
    2. Build a horizontal **density profile** across the block bbox:
       ``profile[x] = number of lines that have any word covering x``.
    3. A "white corridor" is a contiguous x-range where the profile is
       ≤ ``ceil(20 % × n_lines)`` AND its width ≥ ``0.3 inch``. Two
       corridors separated by a thin density spike (typically a stray
       noise word) are merged.
    4. Each corridor is a column **boundary**. ``len(corridors) + 1``
       columns are produced. Words from the *original* (un-filtered)
       lines are then assigned to columns by *centre* x.
    5. The block is rejected as not-a-table when no corridors exist
       (so a multi-line paragraph with full-width text falls back to
       a normal text region) or when fewer than 60 % of resulting
       rows have ≥ 2 populated cells.

    The corridor approach is much more robust than cluster-based
    column detection: a paragraph's word-start positions form many
    accidental clusters, but a paragraph never has wide vertical
    white strips because the text fills the line width. Tables, on
    the other hand, *always* have wide vertical white strips between
    columns.
    """
    raw_lines = block.get("lines") or []
    if len(raw_lines) < 3:
        return None  # need a header + ≥ 2 data rows to count as a table

    bbox_l, _, bbox_r, _ = block["bbox"]
    width = bbox_r - bbox_l
    if width <= 0:
        return None

    # Step 1: drop obvious OCR border noise so it doesn't punch holes
    # in our column corridors.
    min_word_w = max(8, dpi // 12)

    def _is_noise(w: dict) -> bool:
        text = w["text"].strip()
        if not text:
            return True
        if len(text) == 1 and not text.isalnum():
            return True
        if w["width"] < min_word_w:
            return True
        return False

    clean_lines: list[dict] = []
    for line in raw_lines:
        cleaned = [w for w in line["words"] if not _is_noise(w)]
        if cleaned:
            clean_lines.append({**line, "words": cleaned})
    if len(clean_lines) < 3:
        return None

    n_lines = len(clean_lines)

    # Step 2: horizontal density profile (per-line presence, not per-word).
    profile = [0] * (width + 1)
    for line in clean_lines:
        present = [False] * (width + 1)
        for w in line["words"]:
            l = max(0, w["left"] - bbox_l)
            r = min(width, w["left"] + w["width"] - bbox_l)
            for x in range(l, r + 1):
                present[x] = True
        for x in range(width + 1):
            if present[x]:
                profile[x] += 1

    # Step 3: locate white corridors.
    # density_threshold = ceil(20% of lines) so a stray word on one
    # line doesn't punch a corridor; min_gap_px = 0.35 inch is wider
    # than any inter-word gap and slightly wider than typical "wide
    # word" gaps (so an isolated wide word inside the gutter does not
    # split a real column boundary in two).
    density_threshold = max(0, n_lines // 5)
    min_gap_px = max(50, int(dpi * 0.35))

    corridors: list[tuple[int, int]] = []
    in_gap = False
    gap_start = 0
    for x in range(width + 1):
        if profile[x] <= density_threshold:
            if not in_gap:
                in_gap = True
                gap_start = x
        else:
            if in_gap:
                in_gap = False
                if x - gap_start >= min_gap_px:
                    corridors.append((gap_start, x))
    if in_gap and (width + 1 - gap_start) >= min_gap_px:
        corridors.append((gap_start, width + 1))

    # Drop the leading + trailing corridors (those are the left/right
    # page margins of the block, not column dividers).
    if corridors and corridors[0][0] == 0:
        corridors = corridors[1:]
    if corridors and corridors[-1][1] >= width:
        corridors = corridors[:-1]

    # Merge corridors when the "spike" between them is sparse (i.e. it
    # represents stray noise such as a misread border or an isolated
    # word from one row), but KEEP them separate when the spike is
    # dense (i.e. it represents real column content where most rows
    # have a word). The decisive signal is the spike's *peak density*,
    # not its width — a real narrow column (e.g. a single-digit
    # quantity in an invoice table) can be only ~50 px wide but every
    # row has a word in it; a noise spike has at most 1–2 rows with
    # any content.
    merged: list[tuple[int, int]] = []
    half_lines = max(2, n_lines // 2)
    for c in corridors:
        if not merged:
            merged.append(c)
            continue
        spike_l = merged[-1][1]
        spike_r = c[0]
        if spike_r <= spike_l:
            merged.append(c)
            continue
        spike_peak = max(profile[spike_l:spike_r], default=0)
        if spike_peak < half_lines:
            # Sparse spike → noise → merge corridors.
            merged[-1] = (merged[-1][0], c[1])
        else:
            # Dense spike → real column → keep corridors separate.
            merged.append(c)
    corridors = merged

    if not corridors:
        return None  # no column dividers → not a table

    # Step 4: column dividers — use the *midpoint* of each corridor as
    # the boundary between adjacent columns. This makes the column
    # ranges contiguous so a word whose centre falls anywhere from one
    # column's nominal start to the start of the next column gets
    # assigned correctly. Using the corridor *start* would leave the
    # corridor itself in no column at all and any word whose bbox
    # extended even slightly into the corridor (e.g. a wide header
    # like "Quantity") would mis-fall to the wrong column.
    dividers = [bbox_l + (cl + cr) // 2 for cl, cr in corridors]
    col_ranges: list[tuple[int, int]] = []
    prev = bbox_l
    for d in dividers:
        col_ranges.append((prev, d))
        prev = d
    col_ranges.append((prev, bbox_r))
    n_cols = len(col_ranges)

    # Build the row × col grid using the ORIGINAL (noise-included)
    # words so we don't lose any cell content. Each word goes into
    # the column whose range contains its centre x.
    def _col_for_word(w: dict) -> int:
        cx = w["left"] + w["width"] / 2
        for i, (cl, cr) in enumerate(col_ranges):
            if cl <= cx < cr:
                return i
        return n_cols - 1

    rows: list[list[str]] = []
    for line in raw_lines:
        row = [""] * n_cols
        for w in sorted(line["words"], key=lambda w: w["left"]):
            ci = _col_for_word(w)
            row[ci] = (row[ci] + " " + w["text"]).strip() if row[ci] else w["text"]
        rows.append(row)

    # Step 5: sanity check — at least 60 % of rows must have words in
    # ≥ 2 distinct columns. Otherwise we caught a structural false
    # positive (e.g. a list of left-aligned items where one happened
    # to wrap to a "second column").
    n_multi = sum(1 for r in rows if sum(1 for c in r if c.strip()) >= 2)
    if n_multi / len(rows) < 0.6:
        return None

    return {
        "type": "table",
        "bbox": block["bbox"],
        "rows": rows,
    }


def _detect_image_regions(
    png_path: Path,
    word_boxes: list[tuple[int, int, int, int]],
    dpi: int,
) -> list[dict]:
    """Detect figure regions on a rendered page bitmap.

    Algorithm:

    1. Load the page PNG, convert to grayscale, binarize (anything
       darker than ~200/255 is "content").
    2. Mask out tesseract's word bounding boxes (paint them out) so
       the remaining content is non-text — figures, photos, charts,
       borders, etc.
    3. Dilate to merge neighbouring blobs into single figure regions.
    4. Find external contours, filter by minimum area + minimum
       dimension to drop noise / sliver artifacts.
    5. For each surviving region, crop the **color** page bitmap and
       PNG-encode the crop as bytes.

    Returns a list of ``{"type": "image", "bbox", "png_bytes"}`` dicts
    in the same shape as text/columns regions, ready to merge with
    the layout from ``_build_regions``.

    OpenCV is already a transitive dependency via pdf2docx
    (opencv-python-headless), so this adds no new install requirement.
    """
    try:
        import cv2  # noqa: PLC0415
        import numpy as np  # noqa: PLC0415
    except ImportError:
        return []

    img = cv2.imread(str(png_path), cv2.IMREAD_COLOR)
    if img is None:
        return []
    h, w = img.shape[:2]

    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    _, binary = cv2.threshold(gray, 200, 255, cv2.THRESH_BINARY_INV)

    # Mask out text — pad each word box slightly so anti-aliased glyph
    # edges don't bleed through and create stray contours.
    text_mask = np.zeros_like(binary)
    pad = max(2, dpi // 60)  # ≈5 px at 300 DPI
    for (l, t, r, b) in word_boxes:
        l = max(0, l - pad)
        t = max(0, t - pad)
        r = min(w, r + pad)
        b = min(h, b + pad)
        text_mask[t:b, l:r] = 255
    non_text = cv2.bitwise_and(binary, cv2.bitwise_not(text_mask))

    # Morphological opening: erode then dilate with a small kernel.
    # This removes isolated single-pixel speckles (sensor noise from
    # _simulate_scan, JPEG block artifacts, dust specks from a real
    # scan) while preserving any feature larger than ~3 px in either
    # dimension. Critically this prevents the noise specks scattered
    # in the *page margin* between the photo and the page edge from
    # extending the photo's tight-crop bbox out into the white margin
    # — the root cause of the "thick border around extracted images"
    # the user reported.
    open_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (3, 3))
    denoised = cv2.morphologyEx(non_text, cv2.MORPH_OPEN, open_kernel)

    # Dilate just enough to merge adjacent figure blobs into single
    # connected components. We use a small kernel (∝ DPI) and only
    # ONE iteration so the resulting contour bbox is as close to the
    # real figure edges as possible — extra dilation means extra
    # white-page padding around every cropped image.
    k = max(3, dpi // 40)
    kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (k, k))
    dilated = cv2.dilate(denoised, kernel, iterations=1)

    contours, _ = cv2.findContours(
        dilated, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE,
    )

    # Sanity thresholds, scaled to DPI:
    # * minimum area: half a square inch
    # * minimum dimension: 0.4 inch in either direction
    # * maximum area: 85% of page (don't grab the whole page)
    # * aspect ratio: 0.2 .. 5.0 (rules out long thin slivers)
    # * content density: ≥ 12 % dark pixels in the *original* binary
    #   image, computed inside the candidate bbox. Photos and charts
    #   typically run 25–50 %; table borders run 3–8 % because they
    #   are mostly white interior with thin grid lines, so this
    #   threshold rejects bordered tables that survived the text mask.
    min_area = int((dpi * 0.7) ** 2)
    min_dim = int(dpi * 0.4)
    max_area = int(w * h * 0.85)
    min_density = 0.12

    regions: list[dict] = []
    for cnt in contours:
        x, y, cw, ch = cv2.boundingRect(cnt)
        area = cw * ch
        if area < min_area or area > max_area:
            continue
        if cw < min_dim or ch < min_dim:
            continue
        ratio = cw / ch
        if ratio < 0.2 or ratio > 5.0:
            continue
        # Density check on the *denoised* binary. Using denoised
        # rather than the raw non_text means scattered noise specks
        # don't artificially inflate the density of an empty region.
        roi = denoised[y:y + ch, x:x + cw]
        density = float(roi.sum()) / (255.0 * area)
        if density < min_density:
            continue
        # Tighten the crop to the *actual* content extent. The
        # dilated bounding rect always overshoots the real figure
        # edges (that's how dilation works), and any overshoot ends
        # up as white-page padding around the photo — visually a
        # thick border in the final DOCX. Use the *denoised* mask so
        # isolated noise pixels in the page margin don't pull the
        # tight bbox out into the whitespace.
        if int(roi.sum()) == 0:
            continue
        # Tighten by *per-row/col density*. Using min/max of content
        # pixels lands the bbox on the first row that has any content
        # at all — but for a photo with anti-aliased fade-out edges,
        # that row may be 95 % white background with a single content
        # pixel, and the result is a visible light haze around the
        # extracted photo. Walking inward from each side until the
        # row/column has ≥ 5 % content pixels trims those fade rows
        # without cutting deep into the photo body. The 5 % threshold
        # is conservative — even a photo of mostly-empty sky has 5 %
        # content pixels per edge row.
        density_threshold_frac = 0.05
        col_min = max(1, int(ch * density_threshold_frac))
        row_min = max(1, int(cw * density_threshold_frac))
        # Per-column content count (sum across rows / 255)
        col_counts = (roi > 0).sum(axis=0)
        row_counts = (roi > 0).sum(axis=1)

        cols_with = np.where(col_counts >= col_min)[0]
        rows_with = np.where(row_counts >= row_min)[0]
        if cols_with.size == 0 or rows_with.size == 0:
            continue
        tight_l = max(0, x + int(cols_with[0]))
        tight_t = max(0, y + int(rows_with[0]))
        tight_r = min(w, x + int(cols_with[-1]) + 1)
        tight_b = min(h, y + int(rows_with[-1]) + 1)
        crop = img[tight_t:tight_b, tight_l:tight_r]
        ok, encoded = cv2.imencode(".png", crop)
        if not ok:
            continue
        regions.append({
            "type": "image",
            "bbox": (tight_l, tight_t, tight_r, tight_b),
            "png_bytes": encoded.tobytes(),
        })
    return regions


def _build_regions(
    blocks: list[dict],
    page_width_px: int,
    image_regions: list[dict] | None = None,
    dpi: int = 300,
) -> list[dict]:
    """Group tesseract blocks into a sequence of layout regions.

    Algorithm:

    1. Split blocks into **narrow** (width < ``NARROW_FRAC × page_width``,
       default 55 %) and **wide** (everything else: page-spanning
       headings, captions, full-width paragraphs).

    2. Cluster narrow blocks into vertical **tracks** by left-edge
       x-position. Two narrow blocks belong to the same track when
       their left-edges sit within ``TRACK_TOLERANCE`` (≈ 8 % of page
       width). Each track is then sorted top-to-bottom.

    3. Group tracks whose y-extents overlap (union-find with a
       40 %-of-smaller-extent overlap threshold). A group with ≥ 2
       tracks is a multi-column region; each track becomes one cell of
       the resulting borderless docx table.

    4. Wide blocks and lone narrow blocks become single ``"text"``
       regions in their natural reading order.

    Returns regions sorted top-to-bottom by their bbox y-coordinate so
    the renderer can walk them in reading order.
    """
    if not blocks or page_width_px <= 0:
        return [
            {
                "type": "text",
                "bbox": b["bbox"],
                "block_num": b["block_num"],
                "paragraphs": b["paragraphs"],
            }
            for b in blocks
        ]

    NARROW_FRAC = 0.55
    TRACK_TOLERANCE = page_width_px * 0.08

    # Step 1: split narrow vs wide
    narrow: list[tuple[int, dict]] = []
    for i, b in enumerate(blocks):
        l, _t, r, _btm = b["bbox"]
        if (r - l) < page_width_px * NARROW_FRAC:
            narrow.append((i, b))

    # Step 1b: drop narrow blocks that have NO horizontal sibling at a
    # similar y. A short section heading or figure caption gets a
    # narrow bbox too, but it's not part of any multi-column run
    # because no other block sits next to it on the same horizontal
    # band. Without this filter the heading right above the columns
    # gets sucked into the left column track.
    def _has_horiz_sibling(target: dict) -> bool:
        t_l, t_t, t_r, t_b = target["bbox"]
        t_h = t_b - t_t
        if t_h <= 0:
            return False
        for _, other in narrow:
            if other is target:
                continue
            o_l, o_t, o_r, o_b = other["bbox"]
            v_overlap = min(t_b, o_b) - max(t_t, o_t)
            smaller_h = min(t_h, o_b - o_t)
            if smaller_h <= 0 or v_overlap / smaller_h < 0.3:
                continue
            # Horizontally disjoint (no x-overlap) means it's a sibling
            # in a different column, not part of the same column.
            if o_r <= t_l or o_l >= t_r:
                return True
        return False

    narrow = [(idx, b) for idx, b in narrow if _has_horiz_sibling(b)]

    # Step 2: cluster surviving narrow blocks into tracks by left-edge x.
    narrow.sort(key=lambda x: x[1]["bbox"][0])
    tracks: list[list[tuple[int, dict]]] = []
    for idx, b in narrow:
        bl = b["bbox"][0]
        placed = False
        for track in tracks:
            track_left = sum(t[1]["bbox"][0] for t in track) / len(track)
            if abs(bl - track_left) < TRACK_TOLERANCE:
                track.append((idx, b))
                placed = True
                break
        if not placed:
            tracks.append([(idx, b)])
    # Sort each track top-to-bottom for the cell-render order.
    for track in tracks:
        track.sort(key=lambda x: x[1]["bbox"][1])

    # Step 3: union tracks whose vertical extents overlap. Two tracks
    # form a multi-column group when overlap ≥ 40 % of the smaller
    # track's height. Use a simple union-find.
    n_tracks = len(tracks)
    parent = list(range(n_tracks))

    def _find(i: int) -> int:
        while parent[i] != i:
            parent[i] = parent[parent[i]]
            i = parent[i]
        return i

    def _union(i: int, j: int) -> None:
        ri, rj = _find(i), _find(j)
        if ri != rj:
            parent[ri] = rj

    extents: list[tuple[int, int]] = []
    for track in tracks:
        ts = [b["bbox"][1] for _, b in track]
        bs = [b["bbox"][3] for _, b in track]
        extents.append((min(ts), max(bs)))

    for i in range(n_tracks):
        for j in range(i + 1, n_tracks):
            ti, bi = extents[i]
            tj, bj = extents[j]
            overlap = max(0, min(bi, bj) - max(ti, tj))
            smaller = min(bi - ti, bj - tj)
            if smaller > 0 and overlap / smaller > 0.4:
                _union(i, j)

    track_groups: dict[int, list[int]] = {}
    for i in range(n_tracks):
        track_groups.setdefault(_find(i), []).append(i)

    # Step 4: materialize regions
    regions: list[dict] = []
    used_block_idx: set[int] = set()

    for track_indices in track_groups.values():
        if len(track_indices) < 2:
            continue
        group_tracks = [tracks[i] for i in track_indices]
        # Sort tracks left-to-right by their average left-edge.
        group_tracks.sort(
            key=lambda tr: sum(b["bbox"][0] for _, b in tr) / len(tr),
        )
        for track in group_tracks:
            for idx, _b in track:
                used_block_idx.add(idx)
        all_blocks = [b for tr in group_tracks for _, b in tr]
        bbox = (
            min(b["bbox"][0] for b in all_blocks),
            min(b["bbox"][1] for b in all_blocks),
            max(b["bbox"][2] for b in all_blocks),
            max(b["bbox"][3] for b in all_blocks),
        )
        columns_payload = []
        for track in group_tracks:
            track_blocks = [b for _, b in track]
            columns_payload.append({
                "block_num": None,
                "bbox": (
                    min(b["bbox"][0] for b in track_blocks),
                    min(b["bbox"][1] for b in track_blocks),
                    max(b["bbox"][2] for b in track_blocks),
                    max(b["bbox"][3] for b in track_blocks),
                ),
                "paragraphs": [p for b in track_blocks for p in b["paragraphs"]],
            })
        regions.append({
            "type": "columns",
            "bbox": bbox,
            "columns": columns_payload,
        })

    # Single-block regions: try table detection first; if the block
    # is not tabular, fall back to a "text" region holding its
    # paragraphs.
    for i, b in enumerate(blocks):
        if i in used_block_idx:
            continue
        table = _detect_table_in_block(b, dpi=dpi)
        if table is not None:
            regions.append(table)
            continue
        regions.append({
            "type": "text",
            "bbox": b["bbox"],
            "block_num": b["block_num"],
            "paragraphs": b["paragraphs"],
        })

    # Merge image regions into the stream. Drop any image whose bbox
    # sits entirely inside a text region's bbox — that's almost always
    # a false positive (e.g. table grid lines, the table cells were
    # masked out as text but their borders survived).
    for img in (image_regions or []):
        ix0, iy0, ix1, iy1 = img["bbox"]
        contained = False
        for r in regions:
            if r["type"] != "text":
                continue
            tx0, ty0, tx1, ty1 = r["bbox"]
            if tx0 <= ix0 and ty0 <= iy0 and tx1 >= ix1 and ty1 >= iy1:
                contained = True
                break
        if not contained:
            regions.append(img)

    regions.sort(key=lambda r: r["bbox"][1])
    return regions


def _looks_like_code(text: str) -> bool:
    """Heuristic: does this line look like code, CLI invocation, or JSON?

    Score-based detector keyed on markers that almost never appear in
    natural prose:

    * ``://`` (URL — strong signal)
    * `` -- `` (CLI flag separator)
    * leading/trailing ``{`` or ``}`` (JSON brace on its own line)
    * ``": "`` or ``":"`` (JSON key/value pair)
    * 4+ double-quote characters (densely quoted strings)
    * ratio of structural punctuation ``{}[]()<>=;:|&"'`` > 10%

    Threshold of score ≥ 2 catches CLI commands, single-line JSON, and
    multi-line JSON fragments while leaving body text alone (which
    typically scores 0 or 1 from incidental punctuation).
    """
    if not text:
        return False

    score = 0
    if "://" in text:
        score += 2
    if " -- " in text:
        score += 2
    stripped = text.strip()
    if stripped.startswith("{") or stripped.endswith("{"):
        score += 2
    if stripped.startswith("}") or stripped.endswith("}"):
        score += 2
    if '": "' in text or '":"' in text:
        score += 2
    if text.count('"') >= 4:
        score += 1

    structural = sum(1 for c in text if c in '{}[]()<>=;:|&"\'`')
    if len(text) > 0 and structural / len(text) > 0.10:
        score += 1

    return score >= 2


def _merge_code_paragraphs(paragraphs: list[dict]) -> list[dict]:
    """Mark code-style paragraphs and merge runs of consecutive code
    paragraphs into a single multi-line block.

    Tesseract often splits a JSON code block across several
    "paragraphs" because braces on their own lines have very small
    measured heights compared to the lines with text content, and the
    line-merge step rejects them as too dissimilar. This pass undoes
    that fragmentation by re-merging anything the code detector
    classifies as code.

    Each merged code block:
    * Has ``is_code = True`` so the renderer can switch to monospace.
    * Has ``text`` joined with ``\\n`` so the renderer can preserve line
      breaks via ``run.add_break()``.
    * Inherits ``space_before_pt`` from the first paragraph in the run.
    * Uses the median font size of its constituent lines (smaller and
      more representative than the max).
    """
    for p in paragraphs:
        p["is_code"] = _looks_like_code(p.get("text", ""))

    merged: list[dict] = []
    i = 0
    while i < len(paragraphs):
        p = paragraphs[i]
        if not p.get("is_code"):
            merged.append(p)
            i += 1
            continue

        run_start = i
        while i < len(paragraphs) and paragraphs[i].get("is_code"):
            i += 1
        run_end = i

        block = paragraphs[run_start:run_end]
        if len(block) == 1:
            merged.append(block[0])
            continue

        joined_text = "\n".join(b["text"] for b in block)
        sizes = sorted(
            b["font_size_pt"] for b in block if b.get("font_size_pt")
        )
        font_size = sizes[len(sizes) // 2] if sizes else 0
        merged.append({
            "text": joined_text,
            "font_size_pt": font_size,
            "alignment": "left",
            "space_before_pt": block[0].get("space_before_pt", 0),
            "is_code": True,
        })

    return merged


def _classify_headings(paragraphs: list[dict]) -> None:
    """Tag each paragraph dict with a ``heading_level`` (0..6).

    Heuristic: the body font size is the **statistical mode** (most
    common) across all paragraphs. Anything significantly larger than
    that — by ≥ 30% — is treated as a heading. Heading sizes are then
    clustered (within 1.5pt of each other) and assigned levels in
    descending size order: largest cluster → Heading 1, next → Heading 2,
    capped at H6.

    The function mutates the input list in place. Page-break markers and
    blank-page placeholders are skipped.

    Why heading levels matter: applying Word's built-in ``Heading 1``…
    ``Heading 6`` styles isn't just visual — it tells Word the paragraph
    is *semantically* a heading, which is what populates the navigation
    pane, drives table-of-contents generation, and exposes structure to
    accessibility tools. Direct font-size formatting alone produces a
    document that *looks* structured but is one flat ``Normal`` blob in
    Word's eyes.
    """
    sizes = [
        round(p["font_size_pt"], 1)
        for p in paragraphs
        if p.get("font_size_pt") and not p.get("_page_break")
    ]
    if not sizes:
        return

    # Body size = most common (rounded to 0.5pt for stable mode detection)
    rounded = [round(s * 2) / 2 for s in sizes]
    body_size = Counter(rounded).most_common(1)[0][0]
    threshold = body_size * 1.3

    # Cluster heading-candidate sizes by ≤ 1.5pt proximity, descending.
    distinct = sorted({s for s in sizes if s > threshold}, reverse=True)
    clusters: list[list[float]] = []
    for s in distinct:
        if clusters and clusters[-1][-1] - s <= 1.5:
            clusters[-1].append(s)
        else:
            clusters.append([s])

    size_to_level: dict[float, int] = {}
    for level, cluster in enumerate(clusters[:6], start=1):
        for s in cluster:
            size_to_level[s] = level

    for p in paragraphs:
        if p.get("_page_break") or not p.get("font_size_pt"):
            continue
        s = round(p["font_size_pt"], 1)
        p["heading_level"] = size_to_level.get(s, 0)


def _promote_section_headings(paragraphs: list[dict], gap_threshold_pt: float = 20.0) -> None:
    """Promote any heading paragraph with a large vertical gap above it to
    Heading 1.

    Tesseract's word-height measurements are noisy enough that two
    visually equivalent section headings can end up classified as
    different heading levels (e.g. 23.5pt → H2 and 18pt → H3 even though
    both are top-level sections in the source). The structural signal
    that distinguishes a top-level section heading from a sub-heading is
    *whitespace*: top-level sections have a large gap before them, while
    sub-headings sit close under their parent. We use the
    ``space_before_pt`` already computed by ``_parse_tesseract_tsv`` as
    that signal: any heading with > ``gap_threshold_pt`` of whitespace
    above it is promoted to Heading 1.

    Mutates ``paragraphs`` in place.
    """
    for p in paragraphs:
        if p.get("_page_break"):
            continue
        if not p.get("heading_level"):
            continue
        sb = p.get("space_before_pt", 0) or 0
        if sb > gap_threshold_pt:
            p["heading_level"] = 1


def _populate_paragraph(para, p: dict) -> None:
    """Populate a python-docx paragraph from a parsed paragraph dict.

    Applies heading style, alignment, space-before, and creates runs
    from ``p["text"]`` (with soft line breaks for multi-line code
    blocks). Used by both the linear text emit path and the
    multi-column / table cell emit paths so styling stays consistent.
    """
    level = p.get("heading_level", 0)
    if level > 0:
        # Word's built-in heading styles are named "Heading 1"..6.
        try:
            para.style = para.part.document.styles[f"Heading {level}"]
        except (KeyError, AttributeError):
            pass
    if p.get("alignment") == "center":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if p.get("space_before_pt"):
        para.paragraph_format.space_before = Pt(p["space_before_pt"])

    is_code = p.get("is_code", False)
    lines = p.get("text", "").split("\n")
    prev_run = None
    for line_idx, line_text in enumerate(lines):
        if line_idx > 0 and prev_run is not None:
            prev_run.add_break()
        run = para.add_run(line_text)
        if is_code:
            run.font.name = "Courier New"
        if p.get("font_size_pt"):
            run.font.size = Pt(p["font_size_pt"])
        prev_run = run


def _emit_paragraph(doc: Document, p: dict) -> None:
    """Append a paragraph to the document body."""
    if p.get("_blank_page"):
        doc.add_paragraph()
        return
    if not p.get("text", "").strip():
        return
    para = doc.add_paragraph()
    _populate_paragraph(para, p)


def _emit_table_region(doc: Document, region: dict) -> None:
    """Emit a detected OCR table as a real ``<w:tbl>`` in the DOCX.

    Style choices, mirroring the source PDF's table look:
        * Built-in ``Table Grid`` style → 0.5 pt borders on every cell
        * First row gets a gray fill (#D9D9D9), bold text, centred
          horizontally, and the ``<w:tblHeader/>`` marker so Word
          repeats it on every page if the table breaks across pages
        * Empty cells are preserved so the grid stays aligned even
          when OCR missed a cell.
    """
    rows = region.get("rows") or []
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    if n_cols == 0:
        return
    table = doc.add_table(rows=len(rows), cols=n_cols)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for r_i, row in enumerate(rows):
        for c_i in range(n_cols):
            cell_text = row[c_i] if c_i < len(row) else ""
            cell = table.cell(r_i, c_i)
            cell.text = cell_text
            if r_i == 0:
                _set_cell_shading(cell, "D9D9D9")
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.bold = True
        if r_i == 0 and table.rows:
            _mark_as_header_row(table.rows[0]._tr)


def _emit_image_region(doc: Document, region: dict, dpi: int) -> None:
    """Append a detected figure region as an inline picture.

    Width is computed from the source bbox in pixels divided by the
    OCR DPI to get inches, then capped to 95 % of the page text width
    so the picture never overflows past the page margins.
    """
    from io import BytesIO  # noqa: PLC0415

    bbox_w_px = region["bbox"][2] - region["bbox"][0]
    if bbox_w_px <= 0:
        return
    width_in = bbox_w_px / float(dpi)
    section = doc.sections[0]
    avail = section.page_width - section.left_margin - section.right_margin
    avail_in = avail / 914400.0  # EMU → inches
    if avail_in > 0:
        width_in = min(width_in, avail_in * 0.95)
    doc.add_picture(BytesIO(region["png_bytes"]), width=Inches(width_in))


def _set_section_columns(section, num_cols: int, gutter_twips: int = 720) -> None:
    """Set or replace the ``<w:cols>`` element on a section's sectPr.

    Word measures column gutter (the gap between columns) in twentieths
    of a point — *twips*. 720 twips = 0.5 inch ≈ 1.27 cm, which is the
    Word/LibreOffice default. Without per-column ``<w:col>`` children,
    Word distributes the columns equally — exactly what reference
    two-column documents look like.
    """
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    cols.set(qn("w:num"), str(num_cols))
    cols.set(qn("w:space"), str(gutter_twips))


def _emit_columns_native(doc: Document, region: dict) -> None:
    """Emit a multi-column region using Word's native section columns.

    Word's *Page Layout > Columns* feature creates real newspaper-style
    flow columns via continuous section breaks with ``<w:cols num="N"/>``.
    This is what users mean when they say "two-column layout":

    * Text flows naturally between columns.
    * Layout survives PDF export, viewer round-trips, page breaks.
    * The OOXML matches the user's reference fixture
      ``tests/data/Lorem-ipsum-two-columnts.docx`` byte-for-byte in
      structure (one ``<w:cols num="2"/>`` section followed by a
      continuous break back to ``num="1"``).

    Pattern:

        1. Add a continuous section break and set the NEW section to
           N columns. Content added now belongs to that section.
        2. Emit every column's paragraphs in order, separated by
           explicit column breaks so the OCR'd column boundaries are
           preserved precisely (rather than letting Word auto-balance,
           which would risk putting the wrong paragraph in the wrong
           visual column).
        3. Add another continuous section break with cols=1 to revert
           to single-column for any subsequent full-width content.
    """
    columns = region.get("columns") or []
    n_cols = len(columns)
    if n_cols == 0:
        return

    # Step 1: open a new N-column section.
    multi = doc.add_section(WD_SECTION.CONTINUOUS)
    _set_section_columns(multi, n_cols)

    # Step 2: emit each column's content. Insert a column break at the
    # END of each column (except the last) so the next column's
    # paragraphs land in the next physical column. Attaching the break
    # to the last run of the last paragraph avoids creating an extra
    # blank paragraph at the top of the next column.
    for col_idx, col in enumerate(columns):
        col_paras = [p for p in col["paragraphs"] if p.get("text", "").strip()]
        if not col_paras:
            continue
        for p in col_paras:
            para = doc.add_paragraph()
            _populate_paragraph(para, p)
        if col_idx < n_cols - 1:
            br = OxmlElement("w:br")
            br.set(qn("w:type"), "column")
            last_p = doc.paragraphs[-1]
            if last_p.runs:
                last_p.runs[-1]._element.append(br)
            else:
                last_p.add_run()._element.append(br)

    # Step 3: revert to single-column section for subsequent content.
    tail = doc.add_section(WD_SECTION.CONTINUOUS)
    _set_section_columns(tail, 1)


def ocr_to_docx(
    input_pdf: Path,
    output_docx: Path,
    language: str = "eng",
    dpi: int = 300,
    progress_callback: Optional[ProgressCallback] = None,
) -> None:
    """Build a DOCX directly from a scanned PDF via tesseract OCR.

    Renders each page to PNG via PyMuPDF, runs
    ``tesseract page.png stem -l LANG tsv`` to get word-level text with
    layout grouping, parses the TSV into a structured page layout
    (blocks, multi-column groups, image regions), classifies heading
    levels globally based on font sizes, and writes the result to a
    new DOCX via python-docx with Word's built-in ``Heading 1..6``
    styles applied where appropriate. Multi-column sections are
    rendered as borderless N-column tables; figures detected on the
    page bitmap are inserted as inline pictures. Pages are separated
    by hard page breaks.

    Bypasses pdf2docx entirely — pdf2docx can't process scanned PDFs
    (its layout analyzer requires real glyph metrics, which OCR text
    doesn't have).

    Parameters
    ----------
    input_pdf, output_docx
        Source PDF and destination DOCX.
    language
        Tesseract language code (e.g. ``"eng"``, ``"slk"``, or
        ``"slk+eng"`` for mixed-language documents).
    dpi
        Render resolution; 300 is the tesseract sweet spot for printed
        text. Raise to 400 for fine print, lower to 200 for speed.
    progress_callback
        Optional ``fn(current_page, total_pages)`` invoked once per page
        before OCR is performed.

    Raises
    ------
    RuntimeError
        If tesseract is not installed.
    subprocess.CalledProcessError
        If a tesseract invocation fails (e.g. unknown language code).
    """
    if shutil.which("tesseract") is None:
        raise RuntimeError(
            "tesseract is not installed or not on PATH. "
            "Install it locally first (e.g. `brew install tesseract tesseract-lang`)."
        )

    src = fitz.open(str(input_pdf))
    doc = Document()
    try:
        total = len(src)

        # ---- Pass 1: render + OCR + parse layout per page ----
        # Heading classification needs the whole document's font-size
        # distribution, so we collect everything first and render second.
        page_layouts: list[dict] = []
        all_paragraphs: list[dict] = []

        with tempfile.TemporaryDirectory(prefix="pdf2docx-ocr-") as tmp:
            tmp_dir = Path(tmp)
            for i, page in enumerate(src):
                if progress_callback is not None:
                    progress_callback(i + 1, total)

                pix = page.get_pixmap(dpi=dpi)
                png_path = tmp_dir / f"page-{i:04d}.png"
                pix.save(str(png_path))

                stem = tmp_dir / f"page-{i:04d}"
                subprocess.run(
                    [
                        "tesseract",
                        str(png_path),
                        str(stem),
                        "-l",
                        language,
                        "tsv",
                    ],
                    check=True,
                    capture_output=True,
                )

                tsv_path = stem.with_suffix(".tsv")
                layout = _parse_tesseract_tsv_layout(tsv_path, dpi=dpi)
                image_regions = _detect_image_regions(
                    png_path, layout["word_boxes"], dpi=dpi,
                )
                regions = _build_regions(
                    layout["blocks"],
                    layout["page_width_px"],
                    image_regions=image_regions,
                    dpi=dpi,
                )

                # Collect paragraphs from every region for global
                # heading classification.
                for r in regions:
                    if r["type"] == "text":
                        all_paragraphs.extend(r["paragraphs"])
                    elif r["type"] == "columns":
                        for col in r["columns"]:
                            all_paragraphs.extend(col["paragraphs"])

                page_layouts.append({
                    "regions": regions,
                    "page_width_px": layout["page_width_px"],
                    "page_height_px": layout["page_height_px"],
                })

        # ---- Heading classification across the whole document ----
        _classify_headings(all_paragraphs)
        _promote_section_headings(all_paragraphs)

        # ---- Pass 2: render the structured layout to DOCX ----
        for page_idx, page_layout in enumerate(page_layouts):
            if page_idx > 0:
                doc.add_page_break()
            regions = page_layout["regions"]
            if not regions:
                doc.add_paragraph()  # blank page placeholder
                continue
            for region in regions:
                if region["type"] == "text":
                    for p in region["paragraphs"]:
                        _emit_paragraph(doc, p)
                elif region["type"] == "columns":
                    _emit_columns_native(doc, region)
                elif region["type"] == "table":
                    _emit_table_region(doc, region)
                elif region["type"] == "image":
                    _emit_image_region(doc, region, dpi=dpi)

        doc.save(str(output_docx))
    finally:
        src.close()


def convert_pdf(pdf_path: Path, docx_path: Path, parse_stream_table: bool = True) -> None:
    """Run pdf2docx with sensible defaults."""
    cv = Converter(str(pdf_path))
    try:
        cv.convert(
            str(docx_path),
            start=0,
            end=None,
            # pdf2docx kwargs (see pdf2docx.common.Settings):
            parse_stream_table=parse_stream_table,
        )
    finally:
        cv.close()


def _fix_run_font(run_element) -> bool:
    """If the run uses a 'bad' bullet font, swap it for the OpenType variant.

    Returns True if any change was made.
    """
    rpr = run_element.find(qn("w:rPr"))
    if rpr is None:
        return False
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        return False
    changed = False
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        key = qn(f"w:{attr}")
        val = rfonts.get(key)
        if val in BAD_BULLET_FONTS:
            rfonts.set(key, REPLACEMENT_BULLET_FONT)
            changed = True
    return changed


def fix_bullet_fonts(docx_path: Path) -> int:
    """Walk every run in the document and fix Symbol → SymbolMT.

    Returns the number of runs that were patched.
    """
    doc = Document(str(docx_path))
    patched = 0

    def walk_paragraphs(paragraphs):
        nonlocal patched
        for p in paragraphs:
            for run in p.runs:
                if _fix_run_font(run._element):
                    patched += 1

    walk_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                walk_paragraphs(cell.paragraphs)
                # nested tables
                for nested in cell.tables:
                    for nrow in nested.rows:
                        for ncell in nrow.cells:
                            walk_paragraphs(ncell.paragraphs)

    if patched:
        doc.save(str(docx_path))
    return patched


def _reconstruct_first_table_header(pdf_path: Path) -> tuple[list[str] | None, str | None]:
    """Re-extract the header row of page 1's first table directly from the PDF.

    pdf2docx occasionally drops a table's header row when the header's cell
    structure differs from the data rows (e.g. an empty leftmost cell or
    merged cells below). PyMuPDF's table detector still sees the full table,
    so we use it to recover the header text per logical column plus the
    header background fill color.

    Returns ``(header_texts, fill_hex)`` where ``header_texts`` has one entry
    per logical column. Returns ``(None, None)`` when no usable header band
    can be found.
    """
    doc = fitz.open(str(pdf_path))
    try:
        page = doc[0]
        tabs = page.find_tables()
        if not tabs.tables:
            return None, None
        table = tabs.tables[0]
        table_left, table_top, table_right, _ = table.bbox

        # First data row = the first full-width row that lies below table_top.
        # PyMuPDF over-segments columns when header cells span multiple data
        # columns, but the data row's non-None cells give us the canonical
        # logical column boundaries.
        data_row = None
        for row in table.rows:
            rl, rt, rr, _ = row.bbox
            if (abs(rl - table_left) < 1
                    and abs(rr - table_right) < 1
                    and rt > table_top + 1):
                data_row = row
                break
        if data_row is None:
            return None, None

        # Collect ALL x-edges (both lefts and rights) from non-None data
        # cells, then sort and merge boundaries that are too close
        # together to be real columns.
        #
        # Why this matters: PyMuPDF's `find_tables()` sometimes represents
        # column-divider lines as their own narrow "cells" sitting between
        # the real cells (e.g. 5pt-wide phantom cells between 110pt-wide
        # real columns). The naive "boundary at every cell edge" approach
        # then over-counts columns and the header reconstruction returns
        # more columns than the real table has, which makes the caller
        # refuse to inject the header due to the column-count mismatch.
        #
        # Adaptive tolerance: any pair of boundaries closer than
        # ``max(5pt, 2% of table width)`` is treated as one. 2% of table
        # width scales naturally with the document — for a 700pt-wide
        # table that's 14pt; for a 200pt-wide table that's still at
        # least the 5pt floor. This is below any practical text column
        # width (a 14pt-wide column couldn't fit a single character of
        # 11pt text) and well above the typical phantom-cell width.
        all_xs: list[float] = []
        for cell in data_row.cells:
            if cell is None:
                continue
            all_xs.append(cell[0])
            all_xs.append(cell[2])
        all_xs.sort()

        table_width = table_right - table_left
        merge_tolerance = max(5.0, 0.02 * table_width)

        col_xs: list[float] = []
        for x in all_xs:
            if not col_xs or x - col_xs[-1] > merge_tolerance:
                col_xs.append(x)
        if len(col_xs) < 2:
            return None, None

        header_y_top = table_top
        header_y_bottom = data_row.bbox[1]
        if header_y_bottom - header_y_top < 1:
            return None, None  # no header band above the first data row

        header_texts: list[str] = []
        for i in range(len(col_xs) - 1):
            bbox = (col_xs[i], header_y_top, col_xs[i + 1], header_y_bottom)
            text = page.get_text("text", clip=bbox).strip()
            header_texts.append(" ".join(text.split()))

        if not any(header_texts):
            return None, None  # nothing to inject

        # Sample a filled rectangle inside the header band for the fill color.
        fill_hex: str | None = None
        for d in page.get_drawings():
            fill = d.get("fill")
            rect = d.get("rect")
            if fill is None or rect is None:
                continue
            cx = (rect[0] + rect[2]) / 2
            cy = (rect[1] + rect[3]) / 2
            if (table_left - 1 <= cx <= table_right + 1
                    and header_y_top - 1 <= cy <= header_y_bottom + 1):
                r, g, b = fill[:3]
                fill_hex = (
                    f"{int(round(r * 255)):02X}"
                    f"{int(round(g * 255)):02X}"
                    f"{int(round(b * 255)):02X}"
                )
                break

        return header_texts, fill_hex
    finally:
        doc.close()


def _set_cell_shading(cell, hex_color: str) -> None:
    """Apply a solid background fill to a python-docx cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _mark_as_header_row(tr_element) -> None:
    """Set <w:tblHeader/> on a row so Word treats it as a repeating header."""
    trPr = tr_element.find(qn("w:trPr"))
    if trPr is None:
        trPr = OxmlElement("w:trPr")
        tr_element.insert(0, trPr)
    if trPr.find(qn("w:tblHeader")) is None:
        trPr.append(OxmlElement("w:tblHeader"))


def fix_first_table_header(pdf_path: Path, docx_path: Path) -> bool:
    """Re-inject a missing header row into the first table of the DOCX.

    Detection is conservative: the inserted row is only added when the
    reconstructed header column count matches the existing table's column
    count, and only when the table's current first row is not already that
    header (so re-running the script is safe).

    Also removes the orphaned paragraphs immediately above the table that
    pdf2docx left behind containing the header text fragments.
    """
    header_texts, fill_hex = _reconstruct_first_table_header(pdf_path)
    if header_texts is None:
        return False

    doc = Document(str(docx_path))
    if not doc.tables:
        return False
    table = doc.tables[0]

    if len(table.columns) != len(header_texts):
        return False  # column mismatch — refuse to mangle the table

    def _norm(s: str) -> str:
        return " ".join((s or "").split()).lower()

    existing_first = [_norm(c.text) for c in table.rows[0].cells]
    if existing_first == [_norm(h) for h in header_texts]:
        return False  # already fixed; nothing to do

    # Build the new row by appending then moving its <w:tr> to the top.
    new_row = table.add_row()
    for cell, text in zip(new_row.cells, header_texts):
        cell.text = text
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
        if fill_hex:
            _set_cell_shading(cell, fill_hex)

    tbl_el = table._tbl
    tr_el = new_row._tr
    tbl_el.remove(tr_el)
    first_existing_tr = tbl_el.find(qn("w:tr"))
    first_existing_tr.addprevious(tr_el)
    _mark_as_header_row(tr_el)

    # Remove orphaned paragraphs above the table that contain header text.
    # We walk backwards from the table and only delete paragraphs whose
    # tokens are a strict subset of the reconstructed header tokens — this
    # avoids touching unrelated content above.
    body = doc.element.body
    children = list(body.iterchildren())
    tbl_index = children.index(tbl_el)

    header_tokens: set[str] = set()
    for h in header_texts:
        header_tokens.update(h.lower().split())

    if header_tokens:
        for prev in reversed(children[:tbl_index]):
            if prev.tag != qn("w:p"):
                break
            text = "".join(t.text or "" for t in prev.iter(qn("w:t"))).strip()
            if not text:
                continue  # blank paragraph — skip without breaking the walk
            tokens = set(text.lower().split())
            if tokens and tokens.issubset(header_tokens):
                body.remove(prev)
            else:
                break

    doc.save(str(docx_path))
    return True


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(
        description="Convert PDF to editable DOCX with formatting preserved.",
    )
    parser.add_argument(
        "--version",
        action="version",
        version=(
            f"{__codename__} %(prog)s {__version__} ({__release_date__}) "
            f"— by {__author__}"
        ),
    )
    parser.add_argument("pdf", type=Path, help="Input PDF file")
    parser.add_argument("docx", type=Path, nargs="?", help="Output DOCX file (default: same basename as PDF)")
    # Stream-table detection is OFF by default — for school-curriculum-
    # style PDFs (the dominant use case here), pdf2docx's stream-table
    # heuristic invents tables out of label/value blocks and produces
    # garbage. Opt back in with --stream-tables when you have a PDF that
    # genuinely needs it (e.g. label-less ASCII tables). The GUI uses
    # the same default (the "Disable stream-table detection" checkbox is
    # checked on launch).
    parser.add_argument(
        "--stream-tables",
        action="store_true",
        help="Enable pdf2docx stream-table detection. Off by default — only "
             "use this when you have a PDF that genuinely needs stream-mode "
             "table detection. For most curriculum-style PDFs, leaving it "
             "off gives much cleaner output.",
    )
    parser.add_argument(
        "--ocr",
        action="store_true",
        help="Treat the PDF as scanned: OCR each page with tesseract and "
             "build the DOCX directly from the recognized text. Bypasses "
             "pdf2docx (which can't handle scanned PDFs).",
    )
    parser.add_argument(
        "--ocr-language",
        default="eng",
        help="Tesseract language code for --ocr (default: eng). Use a code "
             "from `tesseract --list-langs`, or combine with '+' (e.g. "
             "'slk+eng') for mixed-language documents.",
    )
    args = parser.parse_args(argv)

    pdf_path: Path = args.pdf
    if not pdf_path.exists():
        print(f"error: {pdf_path} does not exist", file=sys.stderr)
        return 1

    docx_path: Path = args.docx or pdf_path.with_suffix(".docx")

    if args.ocr:
        print(f"OCR-converting {pdf_path} -> {docx_path} (lang={args.ocr_language})")
        def _progress(cur, total):
            print(f"  page {cur}/{total}", flush=True)
        ocr_to_docx(
            pdf_path,
            docx_path,
            language=args.ocr_language,
            progress_callback=_progress,
        )
        print("Done.")
        return 0

    print(f"Converting {pdf_path} -> {docx_path}")
    convert_pdf(pdf_path, docx_path, parse_stream_table=args.stream_tables)
    patched = fix_bullet_fonts(docx_path)
    print(f"Patched {patched} bullet runs (Symbol -> {REPLACEMENT_BULLET_FONT})")
    header_fixed = fix_first_table_header(pdf_path, docx_path)
    print(f"First table header: {'re-injected from PDF' if header_fixed else 'unchanged'}")
    print("Done.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
